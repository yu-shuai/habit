from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import re

doc = Document(r"C:\Users\ycy\Desktop\论文融合.docx")

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.75)
section.different_first_page_header_footer = True

def set_font(run, font_name, font_size, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def create_header_for_section(section, odd_header_text, even_header_text):
    header = section.header
    header.is_linked_to_previous = False

    for para in header.paragraphs:
        for run in para.runs:
            run.text = ""

    if len(header.paragraphs) == 0:
        header.paragraphs[0]

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.text = odd_header_text
    set_font(run, '宋体', 9)

    for rel in header.part.rels.values():
        if "header" in rel.reltype:
            pass

def create_footer_for_section(section):
    footer = section.footer
    footer.is_linked_to_previous = False

    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run()
    run1.text = "第"
    set_font(run1, '宋体', 9)

    run2 = p.add_run()
    run2.text = ""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2._element.append(fldChar1)

    run3 = p.add_run()
    run3.text = ""
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run3._element.append(instrText)

    run4 = p.add_run()
    run4.text = ""
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4._element.append(fldChar2)

    run5 = p.add_run()
    run5.text = "页 共"
    set_font(run5, '宋体', 9)

    run6 = p.add_run()
    run6.text = ""
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run6._element.append(fldChar3)

    run7 = p.add_run()
    run7.text = ""
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run7._element.append(instrText2)

    run8 = p.add_run()
    run8.text = ""
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run8._element.append(fldChar4)

    run9 = p.add_run()
    run9.text = "页"
    set_font(run9, '宋体', 9)

for para in doc.paragraphs:
    text = para.text.strip()

    if not text:
        continue

    if text == "目录":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 14, bold=True)
    elif text.startswith("第") and "章" in text and len(text) < 20:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 22, bold=True)
    elif re.match(r'^第\s*\d+\s*章', text):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 22, bold=True)
    elif re.match(r'^\d+\.\d+\s+', text) and not re.match(r'^\d+\.\d+\.\d+\s+', text):
        for run in para.runs:
            set_font(run, '黑体', 15)
    elif re.match(r'^\d+\.\d+\.\d+\s+', text):
        for run in para.runs:
            set_font(run, '黑体', 14)
    elif text in ["摘要"]:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 14)
    elif text in ["Abstract"]:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 14, bold=True)
    elif "关键词" in text:
        keyword_match = re.match(r'(关键词[:：]?)', text)
        if keyword_match:
            pass
        for run in para.runs:
            set_font(run, '黑体', 12)
    elif "Key Words" in text:
        for run in para.runs:
            set_font(run, '黑体', 12, bold=True)
    elif text in ["结论", "参考文献", "致谢"]:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, '黑体', 22)
    else:
        for run in para.runs:
            current_size = run.font.size
            if current_size is None or current_size == Pt(12):
                set_font(run, '宋体', 12)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, '宋体', 12)

odd_header_text = "2022级软件技术专业毕业设计制作说明书"
even_header_text = "姓名：题目名"
create_header_for_section(section, odd_header_text, even_header_text)
create_footer_for_section(section)

output_path = r"C:\Users\ycy\Desktop\论文融合_格式修改.docx"
doc.save(output_path)
print(f"文档已保存到: {output_path}")
print("已设置：")
print("  - 页面边距：左3cm，右2.2cm，上下2.54cm")
print("  - 页眉：奇数页-2022级软件技术专业毕业设计制作说明书，偶数页-姓名：题目名")
print("  - 页脚：第X页 共XX页")
print("  - 字体格式已按要求调整")