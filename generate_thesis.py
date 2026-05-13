from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path


OUT = Path("C:/Users/ycy/Desktop/\u8bba\u6587.docx")
TMP_OUT = Path("C:/Users/ycy/Desktop/\u8bba\u6587_新版待替换.docx")
PROJECT = Path(r"C:\Users\ycy\Desktop\habit")
TITLE = "基于React和Supabase的移动端习惯养成应用设计与实现"
GRADE_MAJOR = "XX级XX专业毕业设计制作说明书"
STUDENT_LINE = "XXX（姓名）：基于React和Supabase的移动端习惯养成应用设计与实现"


def set_doc_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    settings = doc.settings._element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def set_section_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = False


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*(color if color else (0, 0, 0)))


def add_paragraph(doc, text="", style=None, align=None, font="宋体", size=12, bold=False,
                  first_indent=True, before=0, after=0, line=1.5):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Pt(24) if first_indent else None
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run_font(r, font, size, bold)
    return p


def add_title(doc, text, size, level=None, center=False, before=12, after=8):
    p = add_paragraph(
        doc, text, align=WD_ALIGN_PARAGRAPH.CENTER if center else None,
        font="黑体", size=size, bold=True, first_indent=False, before=before, after=after
    )
    if level:
        p.style = f"Heading {level}"
        for r in p.runs:
            set_run_font(r, "黑体", size, True)
    return p


def add_field(run, instr):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)
    cached = OxmlElement("w:t")
    cached.text = "1"
    run._r.append(cached)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_paragraph_field(paragraph, instr, cached_text="1"):
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)

    r = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    r._r.append(instr_text)

    r = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    r._r.append(fld_separate)

    r = paragraph.add_run(cached_text)
    set_run_font(r, "宋体", 9)

    r = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_end)


def set_headers_footers(section, with_footer=True):
    for header, text in [(section.header, GRADE_MAJOR), (section.even_page_header, STUDENT_LINE)]:
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, "宋体", 9)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not with_footer:
        return
    for text in ["第", "页 共", "页"]:
        pass
    r = p.add_run("第")
    set_run_font(r, "宋体", 9)
    add_paragraph_field(p, " PAGE ", "1")
    r = p.add_run("页 共")
    set_run_font(r, "宋体", 9)
    add_paragraph_field(p, " SECTIONPAGES ", "1")
    r = p.add_run("页")
    set_run_font(r, "宋体", 9)


def restart_page_number(section, start=1):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))


def add_cover(doc):
    add_paragraph(doc, "", first_indent=False, after=80)
    p = add_paragraph(doc, "毕业设计制作说明书", align=WD_ALIGN_PARAGRAPH.CENTER,
                      font="黑体", size=22, bold=True, first_indent=False, after=30)
    p = add_paragraph(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER,
                      font="黑体", size=18, bold=True, first_indent=False, after=80, line=1.8)
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4), Cm(8)]
    items = [("学生姓名", "XXX"), ("学号", "XXXXXXXX"), ("年级专业", "XX级XX专业"),
             ("指导教师", "XXX"), ("完成时间", "2026年5月"), ("设计作品", "Habit 习惯追踪社交 App")]
    for row, item in zip(table.rows, items):
        for i, w in enumerate(widths):
            row.cells[i].width = w
            row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = item[0]
        row.cells[1].text = item[1]
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    set_run_font(r, "宋体", 12)
    add_paragraph(doc, "", first_indent=False)
    add_paragraph(doc, "说明：姓名、学号、年级专业、指导教师可在提交前按学校要求替换。", align=WD_ALIGN_PARAGRAPH.CENTER,
                  font="宋体", size=10, first_indent=False)


def add_abstract(doc):
    doc.add_page_break()
    add_title(doc, "摘要", 14, center=True, before=0)
    paras = [
        "随着移动互联网和云服务的快速发展，用户对个人健康管理、学习计划管理和习惯养成工具提出了更高要求。传统待办类应用多强调任务记录，缺少连续反馈、社交陪伴和团队约束机制，导致用户在长期坚持过程中容易出现动力不足、反馈滞后和数据沉淀价值不明显等问题。针对上述问题，本毕业设计开发了一款名为 Habit 的习惯追踪社交 App。",
        "系统采用 React 19、TypeScript、Vite 和 Tailwind CSS 构建前端界面，使用 Capacitor 将 Web 应用封装为 Android/iOS 移动端应用，后端采用 Supabase 提供用户认证、PostgreSQL 数据库、Storage 文件存储和 Realtime 实时订阅能力。系统围绕个人习惯、团队挑战、打卡动态、好友关注、点赞评论、通知中心、勋章奖励和提醒设置等模块展开设计与实现。",
        "本文首先分析习惯养成类应用的需求背景和业务目标，然后从系统总体架构、功能模块、数据库结构、关键业务流程、前端状态管理、后端安全策略和移动端构建部署等方面进行说明。系统已经完成落地实现，能够支持用户注册登录、创建习惯、每日打卡、上传图片、生成动态、参与团队挑战、处理好友关系以及获得勋章奖励，具备较完整的毕业设计作品形态和实际应用价值。"
    ]
    for t in paras:
        add_paragraph(doc, t)
    p = add_paragraph(doc, first_indent=False, after=12)
    r = p.add_run("关键词：")
    set_run_font(r, "黑体", 14, True)
    r = p.add_run("习惯追踪、移动应用、React、TypeScript、Supabase、Capacitor")
    set_run_font(r, "宋体", 12)

    add_title(doc, "Abstract", 14, center=True, before=18)
    for t in [
        "With the rapid development of mobile Internet and cloud services, users increasingly expect habit-building tools to provide continuous feedback, social encouragement and team-based constraints. Traditional task management applications mainly focus on recording tasks, while they often lack long-term motivation mechanisms and interactive feedback.",
        "This graduation project designs and implements Habit, a mobile social habit tracking application. The front-end is developed with React 19, TypeScript, Vite and Tailwind CSS. Capacitor is used to package the web application into a mobile app, and Supabase provides authentication, PostgreSQL database, storage and realtime subscription capabilities.",
        "The thesis describes the requirement analysis, system architecture, functional modules, database design, key business processes, implementation details, testing and deployment. The application has been implemented and deployed, supporting user authentication, habit creation, daily check-in, image upload, activity feeds, team challenges, friend interaction, notification center and medal rewards."
    ]:
        add_paragraph(doc, t, font="Times New Roman", size=12)
    p = add_paragraph(doc, first_indent=False)
    r = p.add_run("Key Words: ")
    set_run_font(r, "黑体", 16, True)
    r = p.add_run("Habit Tracking; Mobile Application; React; TypeScript; Supabase; Capacitor")
    set_run_font(r, "黑体", 14)


def add_toc(doc):
    doc.add_page_break()
    add_title(doc, "目录", 14, center=True, before=0)
    entries = [
        ("第1章  绪论", 1),
        ("1.1  研究背景", 2), ("1.2  设计目的与意义", 2), ("1.3  国内外现状简述", 2), ("1.4  本文主要工作", 2),
        ("第2章  系统需求分析", 1),
        ("2.1  可行性分析", 2), ("2.2  功能需求分析", 2), ("2.3  非功能需求分析", 2),
        ("第3章  系统总体设计", 1),
        ("3.1  系统架构设计", 2), ("3.2  功能模块设计", 2), ("3.3  业务流程设计", 2),
        ("第4章  数据库设计", 1),
        ("4.1  数据库概念结构", 2), ("4.2  数据表设计", 2), ("4.3  数据安全设计", 2),
        ("第5章  系统详细设计与实现", 1),
        ("5.1  前端界面实现", 2), ("5.2  用户认证与资料模块", 2), ("5.3  习惯与打卡模块", 2), ("5.4  团队挑战模块", 2), ("5.5  社交互动与通知模块", 2),
        ("第6章  系统测试与部署", 1),
        ("6.1  测试环境", 2), ("6.2  功能测试", 2), ("6.3  移动端构建与部署", 2),
        ("结论", 1), ("参考文献", 1), ("附录", 1), ("致谢", 1)
    ]
    for text, level in entries:
        p = add_paragraph(doc, first_indent=False, after=3, line=1.2)
        p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.8)
        r = p.add_run(text)
        set_run_font(r, "黑体", 14 if level <= 2 else 12, True)
        r = p.add_run(" " + "." * max(4, 55 - len(text) * 2) + " ")
        set_run_font(r, "宋体", 12)
        r = p.add_run(" ")
        set_run_font(r, "宋体", 12)


def add_diagram_placeholder(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F3F6FA")
    for p in cell.paragraphs:
        p.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run(title)
    set_run_font(r, "黑体", 12, True, (35, 73, 114))
    for line in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        r = p.add_run(line)
        set_run_font(r, "宋体", 11)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def caption(doc, text):
    add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, font="宋体", size=10.5, first_indent=False, before=4, after=8, line=1.2)


def screenshot_box(doc, caption_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(8)
    shade_cell(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    for _ in range(9):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    r = p.add_run("此处替换为 App 真实截图")
    set_run_font(r, "黑体", 12, True, (120, 120, 120))
    for _ in range(9):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
    caption(doc, caption_text)


def add_table(doc, title, headers, rows, widths=None):
    caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if widths and i < len(widths):
                cell.width = Cm(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.line_spacing = 1.2
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(cell.text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, "宋体", 10.5, row is table.rows[0])
    add_paragraph(doc, "", first_indent=False, after=4)


def add_para_list(doc, paragraphs):
    for text in paragraphs:
        add_paragraph(doc, text)


def add_body(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page(sec)
    restart_page_number(sec, 1)
    set_headers_footers(sec, True)

    add_title(doc, TITLE, 18, center=True, before=0, after=16)

    add_title(doc, "第1章  绪论", 16, level=1, center=True)
    add_title(doc, "1.1  研究背景", 15, level=2)
    for t in [
        "移动互联网已经成为用户日常学习、工作和生活管理的重要入口。习惯养成类产品能够帮助用户把长期目标拆分为每日行为，并通过记录、提醒、反馈等方式提高目标完成率。相比传统纸质记录和简单待办工具，移动端应用能够结合账号体系、云端数据、图片记录、即时反馈和社交关系，为习惯坚持提供更完整的场景支持。",
        "在实际使用中，许多用户并不是缺少目标，而是缺少持续反馈和外部约束。单人习惯容易在新鲜感消退后中断，缺少同伴监督的目标也难以形成长期动力。因此，本设计在基础习惯追踪之外加入好友、关注、动态广场、点赞评论、通知中心和团队挑战，使用户能够在记录自我成长的同时获得来自他人的反馈和陪伴。"
    ]:
        add_paragraph(doc, t)
    add_title(doc, "1.2  设计目的与意义", 15, level=2)
    for t in [
        "本毕业设计的目标是设计并实现一款可落地运行的习惯追踪社交 App，验证 React 前端工程、Supabase 云后端和 Capacitor 移动端封装方案在中小型移动应用中的可行性。系统不仅要求完成用户注册登录、任务创建、每日打卡等基础功能，还需要支持社交互动、团队协作、勋章奖励、移动端部署等综合能力。",
        "从实践意义看，该系统能够为用户提供一种低门槛的习惯记录工具；从技术意义看，项目综合运用了组件化开发、类型约束、状态管理、云数据库、实时订阅、文件存储、移动端打包和业务流程设计等知识，具有较好的毕业设计综合训练价值。"
    ]:
        add_paragraph(doc, t)
    add_title(doc, "1.3  国内外现状简述", 15, level=2)
    for t in [
        "当前市场中已有多种习惯管理和任务管理软件，例如待办清单、打卡记录、日历统计类工具。这些产品在记录与提醒方面较成熟，但部分产品更偏重单人效率管理，社交关系和团队约束较弱；另一些社交产品虽然互动性强，但对目标进度、连续行为和奖励体系的表达不够清晰。",
        "Habit 的设计思路是将习惯追踪、社交反馈和团队挑战结合起来，在轻量移动端界面中呈现任务进度、打卡内容、好友动态和团队协作状态。系统不追求复杂的企业协同，而是面向个人成长场景，强调“每日完成、即时记录、长期坚持、共同监督”。"
    ]:
        add_paragraph(doc, t)
    add_title(doc, "1.4  本文主要工作", 15, level=2)
    for t in [
        "本文主要完成以下工作：第一，分析 Habit App 的业务需求和用户使用场景；第二，设计系统总体架构、功能模块和核心流程；第三，设计以 Supabase PostgreSQL 为核心的数据模型；第四，实现认证、习惯、打卡、动态、好友、关注、通知、团队挑战和勋章等功能；第五，完成移动端构建部署并进行功能测试；第六，总结开发过程中的不足并提出后续优化方向。"
    ]:
        add_paragraph(doc, t)
    add_para_list(doc, [
        "在作品落地过程中，本文没有将系统停留在原型展示层面，而是围绕真实可运行的移动端应用进行实现。前端部分不仅完成页面展示，还处理了用户输入校验、图片压缩预览、乐观更新、下拉刷新、弹窗抽屉管理、移动端触摸交互和主题外观切换；后端部分不仅完成基础表结构，还结合 Supabase Auth、Storage、Realtime 和 Row Level Security 形成完整的数据闭环。",
        "本文写作重点放在“设计依据”和“实现过程”的对应关系上。需求分析部分说明为什么需要这些功能；总体设计部分说明这些功能如何被拆分为模块并组织为系统；数据库设计部分说明数据如何持久化以及如何保证权限边界；详细实现部分说明关键功能在代码层面的实现思路；测试与部署部分说明系统如何验证并最终落地到移动端。"
    ])

    add_title(doc, "第2章  系统需求分析", 16, level=1, center=True)
    add_title(doc, "2.1  可行性分析", 15, level=2)
    add_title(doc, "2.1.1  技术可行性", 14, level=3)
    add_paragraph(doc, "前端采用 React、TypeScript 和 Vite，能够满足组件化界面开发、类型检查和快速构建需求；Tailwind CSS 提供了高效的移动端样式组织方式；Supabase 提供认证、数据库、存储和实时订阅能力，减少了自建后端的复杂度；Capacitor 能够将 Web 应用封装为原生移动应用，适合毕业设计在有限周期内完成完整落地。")
    add_title(doc, "2.1.2  经济可行性", 14, level=3)
    add_paragraph(doc, "系统基于开源前端技术和云服务免费/低成本资源实现，开发环境主要为普通个人电脑、Node.js、Android Studio 和 Supabase 项目，不需要额外采购服务器或商业软件授权，经济成本较低。")
    add_title(doc, "2.1.3  操作可行性", 14, level=3)
    add_paragraph(doc, "系统面向普通移动端用户，底部导航将主要页面划分为首页、好友、任务和我的，交互方式贴近常见 App 使用习惯。用户可通过点击、输入、上传图片、切换标签等方式完成主要操作，学习成本较低。")
    add_title(doc, "2.2  功能需求分析", 15, level=2)
    add_table(doc, "表2-1  系统功能需求表",
              ["模块", "功能需求", "说明"],
              [
                  ["用户认证", "注册、登录、重置密码、修改密码", "基于 Supabase Auth 实现邮箱密码认证"],
                  ["习惯管理", "创建、打卡、删除、归档、继续挑战", "支持个人任务和团队任务"],
                  ["动态记录", "发布文字、图片、可见范围控制", "支持公开、好友、仅自己三种范围"],
                  ["社交互动", "好友申请、关注、点赞、评论、回复", "增强用户长期坚持动力"],
                  ["团队挑战", "邀请码加入、全员打卡、加码投票", "通过团队机制形成外部约束"],
                  ["通知提醒", "互动通知、系统通知、未读数量", "提升反馈及时性"],
                  ["移动端", "构建 Android/iOS 工程", "通过 Capacitor 打包落地"]
              ], [2.8, 5.2, 7.0])
    add_title(doc, "2.2.1  用户角色分析", 14, level=3)
    add_para_list(doc, [
        "系统的主要用户是希望通过移动端记录并坚持长期目标的普通用户。该用户既可以作为个人习惯的创建者，也可以作为团队挑战的队长或成员。个人习惯创建者关注的是任务目标、每日完成状态、断签反馈、历史归档和勋章奖励；团队队长除上述能力外，还需要管理邀请码、开始挑战、移除未开始成员以及在任务完成后发起结算或加码投票；团队成员则主要关注加入团队、查看成员状态、完成每日打卡和参与投票。",
        "从社交关系看，用户还可以作为动态发布者、动态浏览者、好友申请发起者、好友申请接收者、关注者和被关注者。不同角色对应不同的数据权限，例如用户可以编辑自己的资料和动态，但不能删除他人的动态；可以对可见范围允许的动态进行点赞评论，但不能读取他人的私密动态；可以查看发给自己的好友申请，但不能查看无关用户之间的申请记录。"
    ])
    add_title(doc, "2.2.2  典型用例分析", 14, level=3)
    add_table(doc, "表2-2  典型用例说明",
              ["用例", "参与者", "前置条件", "基本流程", "异常处理"],
              [
                  ["注册登录", "未登录用户", "用户拥有邮箱", "输入邮箱密码并提交，认证成功后进入主界面", "邮箱已注册、密码错误或网络异常时给出提示"],
                  ["创建习惯", "登录用户", "用户已登录", "输入名称、目标天数和类型，系统写入 habits 表", "名称为空或数据库写入失败时不创建"],
                  ["发布打卡", "习惯参与者", "存在进行中任务", "选择任务、填写内容、上传图片并发布", "重复打卡时提示或更新当日记录"],
                  ["加入团队", "团队成员", "拥有有效邀请码", "输入邀请码后加入成员列表", "挑战已开始、团队满员或邀请码无效时拒绝"],
                  ["互动通知", "动态浏览者", "动态在可见范围内", "点赞或评论后写入互动数据并生成通知", "RPC 失败时回滚乐观更新"]
              ], [2.2, 2.0, 3.0, 4.6, 4.2])
    add_title(doc, "2.3  非功能需求分析", 15, level=2)
    for t in [
        "易用性方面，系统应具备清晰的页面层级和简洁的移动端交互方式；可靠性方面，打卡记录需要防止重复提交，点赞和评论需要尽量避免并发覆盖；安全性方面，用户数据需要借助 RLS 和认证状态进行隔离；可维护性方面，前端业务逻辑应拆分为独立 Hook，组件保持相对清晰的展示职责；可扩展性方面，团队、通知、勋章等模块应能够在后续迭代中继续扩展。"
    ]:
        add_paragraph(doc, t)
    add_para_list(doc, [
        "性能方面，系统需要适配移动端网络环境。动态列表采用分页加载方式，首次加载限制数量，后续通过加载更多追加数据；图片上传前进行压缩，降低移动网络下的等待时间和存储成本；本地界面采用乐观更新，在数据库写入完成前先给用户即时反馈，避免所有操作都被网络延迟阻塞。",
        "一致性方面，系统需要在用户体验和数据准确性之间取得平衡。打卡操作先检查 habit_logs 表，防止同一用户同一习惯同一天重复打卡；自动动态采用确定性 ID，使同一天同一任务对应同一条记录；点赞和评论采用 RPC 更新 JSONB 数组，减少多个用户同时互动时的覆盖风险；用户修改头像或昵称后，会同步更新自己已发布动态中的用户快照，使历史动态展示保持一致。",
        "可移植性方面，系统以 Web 技术为基础，又通过 Capacitor 打包到移动端。这种方案使大部分业务逻辑能够在浏览器和原生容器中复用，降低了同时维护 Web、Android 和 iOS 多套代码的成本。对于毕业设计而言，该方案能够在较短周期内完成可演示、可安装、可迭代的实际作品。"
    ])

    add_title(doc, "第3章  系统总体设计", 16, level=1, center=True)
    add_title(doc, "3.1  系统架构设计", 15, level=2)
    add_paragraph(doc, "系统采用前后端分离和云服务后端的设计思路。前端 React 应用负责页面渲染、用户交互和部分乐观更新；Supabase 负责认证、数据库读写、文件存储和实时消息推送；Capacitor 负责将 Web 构建产物封装到 Android/iOS 原生工程中。")
    add_para_list(doc, [
        "前端架构中，App.tsx 作为应用组合层，负责初始化 Session、加载基础数据、注册实时订阅并挂载全局弹窗。具体业务逻辑被拆分到多个自定义 Hook 中，例如 useHabitActions 负责习惯创建、打卡、删除、团队操作和奖励结算；useActivityActions 负责动态发布、点赞、评论和可见性修改；useFriendActions 和 useFollowActions 分别处理好友和关注关系；useNotificationActions 处理通知分页、已读状态、声音震动和角标更新。",
        "后端架构中，Supabase Auth 提供认证身份，PostgreSQL 表保存业务数据，Storage bucket 保存头像和打卡图片，Realtime channel 用于监听表变化。前端通过 @supabase/supabase-js 统一访问后端能力，避免单独编写传统 REST 服务。对于点赞和评论这类并发写入场景，系统通过数据库 RPC 进行更新，减少客户端读写回传造成的数据覆盖。",
        "移动端架构中，Vite 将 React 应用构建到 dist 目录，Capacitor 将 dist 作为 webDir 同步进原生工程。Android 工程负责应用包名、启动图、图标、权限和 Gradle 构建，React 应用仍负责主要界面与业务。该结构使系统可以先在浏览器中快速调试，再同步到手机环境测试，开发效率较高。"
    ])
    add_diagram_placeholder(doc, "系统总体架构图", [
        "用户移动端 App",
        "↓ 交互请求 / 实时订阅",
        "React + TypeScript + Vite 前端",
        "↓ Supabase JS Client",
        "Supabase Auth / Database / Storage / Realtime",
        "↓",
        "PostgreSQL 数据表、Storage 文件、Realtime Channel"
    ])
    caption(doc, "图3-1  系统总体架构图")
    add_title(doc, "3.2  功能模块设计", 15, level=2)
    add_diagram_placeholder(doc, "功能模块结构图", [
        "Habit App",
        "├─ 用户认证与资料模块",
        "├─ 习惯任务与每日打卡模块",
        "├─ 团队挑战与投票模块",
        "├─ 动态广场与互动模块",
        "├─ 好友关注与用户主页模块",
        "├─ 通知中心与提醒模块",
        "└─ 设置、外观与移动端更新模块"
    ])
    caption(doc, "图3-2  系统功能模块结构图")
    add_para_list(doc, [
        "用户认证与资料模块是系统入口，负责登录态识别、注册登录、密码重置、资料初始化和资料修改。该模块为其他模块提供当前用户 ID、昵称、头像和自定义 ID 等基础信息。",
        "习惯任务与每日打卡模块是系统核心，负责个人任务和团队任务的创建、展示、进度推进、断签处理、失败归档和奖励结算。该模块与动态模块紧密关联，每次有效打卡都可能生成一条 activities 记录。",
        "社交互动模块负责动态流、好友关系、关注关系、点赞、评论、回复和用户主页。它增强了习惯记录的反馈机制，使用户不仅能看到自己的进度，也能通过好友和关注动态获得外部激励。",
        "通知模块负责将社交行为和系统行为转化为可读消息。例如好友申请、关注、点赞、评论、回复、团队加码投票等操作都可以生成通知。通知既包括数据库中的消息记录，也包括前端声音、震动和未读数量展示。"
    ])
    add_title(doc, "3.3  业务流程设计", 15, level=2)
    add_title(doc, "3.3.1  用户登录流程", 14, level=3)
    add_diagram_placeholder(doc, "登录与资料初始化流程", [
        "打开 App → 检查 Supabase Session",
        "无 Session：展示登录/注册页",
        "有 Session：进入主界面",
        "读取 profiles 表",
        "资料存在：加载用户信息",
        "资料不存在：创建默认资料",
        "加载习惯、动态、好友、关注与通知"
    ])
    caption(doc, "图3-3  用户登录与资料初始化流程图")
    add_title(doc, "3.3.2  打卡流程", 14, level=3)
    add_diagram_placeholder(doc, "每日打卡流程", [
        "选择习惯 → 填写打卡内容 → 上传图片",
        "检查今日是否已有 habit_logs 记录",
        "未打卡：写入 habit_logs",
        "更新 habits 进度与 last_check_date",
        "写入或更新 activities 动态",
        "达到目标：提示结算或继续挑战"
    ])
    caption(doc, "图3-4  每日打卡业务流程图")
    add_title(doc, "3.3.3  团队挑战流程", 14, level=3)
    add_diagram_placeholder(doc, "团队挑战流程", [
        "队长创建团队任务并生成邀请码",
        "成员输入邀请码加入",
        "队长开始挑战并锁定成员",
        "成员每日打卡",
        "全员当日完成后团队进度 +1",
        "达到目标后队长结算或发起加码",
        "全员同意则延长目标，否则结算"
    ])
    caption(doc, "图3-5  团队挑战业务流程图")
    add_title(doc, "3.4  前端状态设计", 15, level=2)
    add_para_list(doc, [
        "系统状态分为三类：第一类是会话和用户资料等基础状态，例如 session、userProfile、friends、followings 和 notifications；第二类是业务列表状态，例如 tasks、completedTasks 和 activities；第三类是界面临时状态，例如当前 Tab、弹窗开关、打卡表单内容、搜索关键字、选中的动态和选中的任务详情。",
        "基础状态和列表状态使用 Zustand 与 React Hook 结合管理，便于不同组件共享；临时界面状态主要保存在 useAppState 中，由 App.tsx 向子组件传递。这样处理的优点是主流程清晰，弹窗、抽屉、Tab 切换等状态集中，缺点是随着功能增加，useAppState 中状态数量较多，后续可继续拆分为 habitStore、socialStore 和 settingsStore。",
        "在用户操作时，系统大量采用乐观更新策略。例如创建任务时先把任务插入本地列表，再写入数据库；点赞时先改变本地 likedBy 数组，再调用 RPC；删除动态时先从本地移除，再向数据库删除。乐观更新可以明显提高移动端交互流畅度，但必须在数据库返回错误时回滚状态，因此每个关键操作都保留了 previous state 或失败回滚逻辑。"
    ])
    add_title(doc, "3.5  异常处理设计", 15, level=2)
    add_para_list(doc, [
        "系统的异常主要分为输入异常、网络异常、权限异常和业务异常。输入异常包括任务名称为空、密码长度不足、搜索关键字为空、邀请码格式错误等，通常在前端直接提示；网络异常包括图片上传失败、数据库写入失败、Realtime 订阅失败等，系统会通过 Toast 提示用户并尽量保留当前可用界面；权限异常由 Supabase RLS 和前端判断共同处理，例如非本人不能编辑他人动态，非队长不能删除进行中的团队任务。",
        "业务异常主要体现在重复打卡、团队满员、挑战已开始、惩罚期断签、加码投票超时等场景。系统在这些场景中不仅要阻止非法操作，还要给出清晰反馈。例如同一习惯当天重复打卡时提示“今日已打卡”；团队挑战开始后邀请码失效，用户不能再加入；团队任务完成后，非队长只能等待队长结算，避免多个成员同时改变团队目标。"
    ])

    add_title(doc, "第4章  数据库设计", 16, level=1, center=True)
    add_title(doc, "4.1  数据库概念结构", 15, level=2)
    add_paragraph(doc, "系统核心数据围绕用户、习惯、打卡记录、动态、好友关系、关注关系、通知和版本发布记录展开。用户通过 profiles 保存资料，通过 habits 保存个人或团队习惯，通过 habit_logs 保存每日打卡记录，通过 activities 沉淀打卡动态和勋章记录。社交关系由 friendships 和 follows 维护，通知由 notifications 表保存。")
    add_para_list(doc, [
        "数据库设计遵循“核心行为独立记录、展示快照适度冗余”的原则。打卡行为通过 habit_logs 独立记录，便于统计用户总打卡天数，也便于检查某天是否已完成；动态表 activities 中保存 user 快照，能够让动态列表在读取时直接显示头像和昵称，减少关联查询次数；当用户更新头像或昵称时，系统再同步更新其历史动态中的 user 字段，以保证展示一致。",
        "团队任务当前将 members 和 vote_status 保存为 JSONB 字段。这样做的原因是团队规模较小，成员状态主要与单个习惯强绑定，使用 JSONB 可以减少表连接并简化快速开发。members 中保存成员 ID、昵称、头像和 lastCheckDate，便于判断团队当天是否全员打卡；vote_status 中保存投票人、选择、加码天数和投票时间，便于判断是否全员同意或是否超时。",
        "对于长期演进，团队成员也可以拆分为 habit_members 表，并将团队投票拆分为 habit_votes 表。关系表方案更适合复杂查询、权限控制和统计分析，但在本毕业设计规模下，JSONB 方案能够更直接地服务于业务闭环。论文中保留了这种取舍说明，体现数据库设计不是简单堆表，而是结合业务规模和开发周期进行权衡。"
    ])
    add_diagram_placeholder(doc, "核心实体关系示意", [
        "profiles 1 ── N habits",
        "profiles 1 ── N habit_logs",
        "habits 1 ── N habit_logs",
        "habits 1 ── N activities",
        "profiles 1 ── N activities",
        "profiles N ── N profiles（friendships / follows）",
        "profiles 1 ── N notifications"
    ])
    caption(doc, "图4-1  数据库核心实体关系图")
    add_title(doc, "4.2  数据表设计", 15, level=2)
    add_table(doc, "表4-1  主要数据表说明",
              ["表名", "主要作用", "关键字段"],
              [
                  ["profiles", "保存用户资料", "id、custom_id、name、avatar"],
                  ["habits", "保存习惯任务", "total_days、current_progress、type、members"],
                  ["habit_logs", "保存每日打卡记录", "habit_id、user_id、completed_date"],
                  ["activities", "保存动态和勋章", "content、images、visibility、liked_by、comments"],
                  ["friendships", "保存好友关系", "requester_id、receiver_id、status"],
                  ["follows", "保存关注关系", "follower_id、following_id"],
                  ["notifications", "保存通知", "type、actor_id、is_read、post_id"],
                  ["app_releases", "保存版本发布信息", "version、build_number、download_url"]
              ], [3.2, 4.5, 7.0])
    add_title(doc, "4.2.1  habits 表设计", 14, level=3)
    add_para_list(doc, [
        "habits 表是系统最重要的业务表之一。对于个人习惯，它记录任务名称、目标天数、当前进度、最后打卡日期、是否归档和是否失败等字段；对于团队习惯，它还记录创建者 ID、邀请码、成员列表、是否开始、投票状态和队长删除标记等字段。通过 type 字段区分 single 和 team，可以在同一张表中统一保存两类任务，减少前端读取任务列表时的复杂度。",
        "current_progress 与 total_days 共同表示任务完成状态。当 current_progress 小于 total_days 时，任务仍处于进行中；当 current_progress 大于或等于 total_days 时，系统进入结算逻辑，用户可以领取奖励或继续挑战。is_archived 用于区分进行中任务和已完成/失败任务，archived_at 用于记录归档时间，便于个人中心展示历史任务。"
    ])
    add_title(doc, "4.2.2  activities 表设计", 14, level=3)
    add_para_list(doc, [
        "activities 表承担动态流和勋章记录两类职责。普通打卡动态的 type 通常为 checkin，保存文字内容、图片数组、习惯标签和可见范围；勋章记录的 type 为 medal，tag 采用 medal:天数 的格式，表示用户获得了某个阶段性奖励。将勋章也作为动态保存，可以让勋章墙直接从 activities 中读取记录，同时避免删除已完成任务时误删用户已获得的勋章。",
        "liked_by 和 comments 采用 JSONB 数组保存，结构中包含用户 ID、昵称、评论内容、创建时间和互动范围。该设计适合轻量社交场景，读取单条动态时可以一次拿到互动数据。为避免并发写入覆盖，前端不直接读出整个数组后覆盖写回，而是调用 add_like、remove_like、add_comment、remove_comment 等 RPC，由数据库侧完成原子修改。"
    ])
    add_title(doc, "4.2.3  notifications 表设计", 14, level=3)
    add_para_list(doc, [
        "notifications 表用于保存通知中心数据。每条通知包含接收人 user_id、触发者 actor_id、触发者昵称和头像快照、通知类型、关联动态、关联评论、通知内容、动态预览、是否已读和创建时间。保存触发者快照可以减少通知列表渲染时的关联查询，即使触发者之后修改昵称，通知也能保持当时语义。",
        "通知类型覆盖 like、comment、reply、friend_request、friend_accept、follow、mention 和 system。系统通知用于团队加码投票、团队成员加入、挑战开始等非普通社交互动场景。前端加载通知时按 created_at 倒序分页，并单独统计 is_read=false 的数量用于未读角标展示。"
    ])
    add_title(doc, "4.3  数据安全设计", 15, level=2)
    add_paragraph(doc, "由于系统直接面向用户数据，数据库安全设计需要重点考虑认证用户的数据隔离。Supabase 通过 Row Level Security 提供行级安全能力，系统在公开 schema 中的业务表应开启 RLS，并根据 auth.uid() 判断当前用户是否有权限读取或写入相关数据。对于点赞、评论等需要非作者更新动态的场景，系统通过 RPC 函数进行原子更新，以减少并发覆盖风险。")
    add_para_list(doc, [
        "在 profiles 表中，用户应只能修改自己的资料，但可以按搜索需求读取有限的公开资料字段，例如昵称、头像和自定义 ID。habits 表应限制普通用户只能读取自己创建或参与的习惯，团队任务可以通过 members 字段判断成员身份。habit_logs 表则应严格限制为本人记录，防止他人伪造打卡记录。",
        "activities 表的权限较复杂。公开动态可以被所有认证用户读取，好友动态需要结合好友关系在应用层过滤，私密动态只能本人读取。动态作者可以编辑或删除自己的动态，其他用户只能通过受控的互动入口进行点赞和评论。为了避免将任意字段更新权限暴露给所有用户，生产环境中更推荐将互动能力封装为 security definer 或受限 RPC，并在函数内部校验可见范围和用户身份。",
        "Storage 安全方面，头像和打卡图片统一保存在 habit bucket 中。路径以用户 ID 开头，例如 userId/avatar.ext 和 userId/posts/timestamp-random.ext。这样的路径设计便于设置策略：认证用户只能上传或覆盖自己目录下的文件，删除动态或注销账号时也能根据 URL 解析路径并清理对应文件。前端只使用 anon key，service_role key 不应出现在任何浏览器环境变量中。"
    ])
    add_title(doc, "4.4  数据一致性与索引设计", 15, level=2)
    add_para_list(doc, [
        "为了保证打卡数据一致性，habit_logs 表应设置唯一约束，约束字段建议为 habit_id、user_id 和 completed_date。这样即使前端由于网络重试或用户重复点击提交，也不会在数据库中形成重复打卡记录。activities 表的自动打卡动态采用 auto-habitId-userId-date 的确定性 ID，同样可以避免同一天同一任务重复生成多条自动动态。",
        "索引设计主要围绕高频查询建立。notifications 表需要按 user_id 和 is_read 查询未读数量，并按 created_at 倒序分页；activities 表需要按 created_at 排序加载动态流，也经常按 user_id 或 habit_id 过滤；friendships 表需要按 requester_id、receiver_id 和 status 查询好友关系；follows 表需要分别按 follower_id 和 following_id 查询关注列表与粉丝列表。",
        "系统当前的数据一致性策略以数据库约束、RPC 原子更新和前端回滚共同完成。数据库负责底线约束，RPC 负责并发数组修改，前端负责在用户交互层面减少重复操作和展示明确反馈。三者结合后，即使在移动端网络不稳定的情况下，也能尽量保持界面状态与云端数据的一致。"
    ])

    add_title(doc, "第5章  系统详细设计与实现", 16, level=1, center=True)
    add_title(doc, "5.1  前端界面实现", 15, level=2)
    add_paragraph(doc, "系统采用移动端优先的界面布局，主界面由顶部 Header、内容区域和底部导航组成。底部导航包含首页、好友、任务和我的四个入口；首页内部又划分为广场、团队、关注三个子标签。界面实现中使用 Tailwind CSS 进行样式组织，使用 lucide-react 提供图标，使用 motion 实现弹窗、列表切换和庆祝动画等过渡效果。")
    add_para_list(doc, [
        "移动端界面设计强调单手操作和层级清晰。顶部 Header 主要放置心情、搜索、创建任务和设置入口，属于全局操作；底部导航放置高频页面入口，用户可以快速切换首页、好友、任务和我的；复杂操作尽量通过抽屉或浮层完成，例如打卡发布、设置、搜索、动态详情、任务详情和勋章详情，避免页面跳转过多造成使用割裂。",
        "组件设计遵循“展示组件与业务 Hook 分离”的思路。HabitCard、MomentItem、BottomNav、Header 等组件主要负责渲染和基础交互；真正涉及数据读写的逻辑放入 useHabitActions、useActivityActions、useFriendActions 等 Hook 中。这样做有利于降低组件复杂度，也便于在后续测试或维护时定位业务逻辑。",
        "界面反馈方面，系统使用 Toast 告知用户操作结果，例如创建失败、今日已打卡、加入团队成功、点赞失败、密码修改成功等。对于加载状态，动态列表和任务列表提供 Skeleton 占位，减少空白等待；对于打卡成功、获得奖励等正向事件，系统使用庆祝动效增强反馈。"
    ])
    screenshot_box(doc, "图5-1  Habit App 首页真实截图（请替换）")
    screenshot_box(doc, "图5-2  任务列表页面真实截图（请替换）")
    add_title(doc, "5.2  用户认证与资料模块", 15, level=2)
    add_paragraph(doc, "认证模块基于 Supabase Auth 实现。用户可使用邮箱和密码注册或登录，也可以通过忘记密码流程发送重置邮件。系统在 App 启动时监听 Supabase Session 状态，未登录时展示 Auth 组件，登录后进入主界面。用户资料由 profiles 表保存，首次登录且资料不存在时，系统会自动创建默认昵称和头像。")
    add_para_list(doc, [
        "注册逻辑需要处理 Supabase 返回的多种情况。若邮箱已注册，系统会识别错误信息并提示用户直接登录；若注册成功且项目开启自动确认，则用户可以直接获得 Session 并进入应用；若需要邮箱确认，则系统提示用户查看邮箱。登录逻辑则主要校验邮箱和密码是否正确，并根据认证状态切换界面。",
        "资料模块的实现包括资料读取、资料创建、昵称修改、自定义 ID 修改和头像上传。fetchProfile 会根据当前 session.user.id 查询 profiles 表；如果没有记录，则创建默认资料。头像上传时，系统先读取本地图片用于预览，再调用 Storage 上传到 habit bucket，上传成功后更新 profiles.avatar。为了避免历史动态显示旧头像，系统还会批量更新 activities 表中当前用户发布动态的 user 快照。",
        "账号安全相关功能包括修改密码、退出登录和注销账号。修改密码调用 supabase.auth.updateUser；退出登录调用 supabase.auth.signOut 并关闭设置弹窗；注销账号在前端删除 public 表中的用户相关数据和 Storage 文件后退出登录。由于前端不能使用 service_role key，彻底删除 auth.users 记录需要通过 Supabase Edge Function 使用 Admin API 实现，这一点在系统设计中作为后续增强说明。"
    ])
    screenshot_box(doc, "图5-3  登录/注册页面真实截图（请替换）")
    add_title(doc, "5.3  习惯与打卡模块", 15, level=2)
    add_paragraph(doc, "习惯模块支持创建个人任务和团队任务。创建任务时，系统生成 Habit 对象并以乐观更新方式先显示到界面，再写入 Supabase 数据库。每日打卡时，系统先查询 habit_logs 表确认当天是否已经打卡，未打卡时写入记录并更新 habits 表中的进度字段。对于个人任务，进度由用户个人打卡推进；对于团队任务，只有所有成员当天均完成打卡时，团队进度才会增加。")
    add_paragraph(doc, "打卡动态支持文字和图片。图片上传前先进行本地压缩和 Base64 预览，上传到 Supabase Storage 后再替换为公共 URL。动态写入 activities 表，并保留点赞和评论数组。系统通过确定性动态 ID 避免同一天同一任务重复产生多条自动动态。")
    add_para_list(doc, [
        "创建任务时，系统根据用户选择的任务类型构造不同数据。个人任务的 is_started 默认为 true，表示创建后即可打卡；团队任务的 is_started 默认为 false，并生成 6 位邀请码，同时将创建者作为第一个成员写入 members 数组。创建动作采用乐观更新：本地先将新任务展示出来，如果 Supabase insert 返回错误，再从本地列表移除该任务并提示失败。",
        "打卡逻辑是系统中条件分支最多的业务之一。首先，系统根据任务 ID 查找当前任务，若任务不存在或已经失败，则直接返回；其次，如果当前进度已经达到目标天数，系统不会继续普通打卡，而是弹出结算逻辑；再次，团队任务若尚未开始，也不能打卡；最后，系统查询 habit_logs 表判断当天是否已存在记录，防止重复打卡。",
        "个人任务的断签惩罚采用两阶段设计。普通状态下，如果用户超过一天未打卡，系统将任务置为 penalty_mode；惩罚状态下，用户需要连续完成 3 天打卡才能解除惩罚并恢复进度增长；如果惩罚期间再次断签，任务被标记为失败并归档。这种设计比简单的“断签即失败”更温和，也更符合习惯养成产品鼓励持续回归的目标。",
        "奖励结算采用 getMedalForDays 方法计算勋章层级。系统根据目标天数向下匹配 7、30、90、180、365、500 等阶段，用户领取奖励后会生成 type 为 medal 的私密动态，并将任务归档。若用户选择继续挑战，系统会提高 total_days 并保留 current_progress，使用户可以在原有成果上继续推进。"
    ])
    screenshot_box(doc, "图5-4  打卡发布页面真实截图（请替换）")
    add_title(doc, "5.4  团队挑战模块", 15, level=2)
    add_paragraph(doc, "团队挑战模块是本系统区别于普通习惯记录工具的重要设计。用户创建团队任务后，系统生成 6 位邀请码，其他用户可在团队页输入邀请码加入。挑战开始前，队长可以移除成员；开始后成员列表锁定，邀请码失效。团队目标完成后，队长可选择直接结算，也可发起加码投票。投票采用全员同意机制，只要存在成员拒绝或 24 小时未完成投票，就按照结算处理。")
    add_para_list(doc, [
        "团队加入流程首先会对邀请码进行 trim 和大写处理，减少输入大小写造成的失败。系统先在本地任务列表中查找对应团队，若不存在则向 Supabase 查询 invite_code。找到团队后，需要依次判断挑战是否已经开始、团队人数是否超过上限、当前用户是否已经在团队中。只有所有条件满足时，才会把当前用户追加到 members 数组并更新数据库。",
        "团队打卡与个人打卡最大的区别在于进度推进条件。成员个人完成打卡后，只更新自己在 members 中的 lastCheckDate；系统随后检查团队中所有成员当天是否都已打卡。如果全员完成，则团队 current_progress 加一；如果仍有成员未完成，则只提示剩余人数。这种机制把个人行为转化为团队共同进度，能够形成更强的监督效果。",
        "团队任务完成后的加码投票采用队长发起、成员表态、全员同意的方式。队长可以选择直接结算，也可以提出更高的新目标天数。提出加码后，vote_status 中保存队长的 proposal，成员在 24 小时内投票。若某个成员选择结算，则视为否决并立即结算；若所有成员都选择继续，则系统更新 total_days 并清空 vote_status；若超时未全员同意，也按照结算处理。",
        "队长删除进行中的团队任务时，系统并不直接从所有成员列表中移除任务，而是将 captain_deleted 标记为 true。这样可以让队员在自己的设备上看到队长已删除的状态，并自行移除相关任务，避免突然消失造成困惑。该设计体现了团队数据在多用户之间同步时需要考虑可解释性。"
    ])
    screenshot_box(doc, "图5-5  团队挑战页面真实截图（请替换）")
    add_title(doc, "5.5  社交互动与通知模块", 15, level=2)
    add_paragraph(doc, "社交模块包括好友、关注、动态、点赞、评论和通知。用户可以通过昵称或自定义 ID 搜索他人，发送好友申请或关注用户。动态互动采用乐观更新策略：用户点赞或评论后，界面先立即更新，再通过 RPC 写入数据库；如果写入失败则回滚本地状态。通知模块根据点赞、评论、回复、好友申请、关注等操作生成通知记录，并通过 Realtime 订阅实时刷新未读数量。")
    add_para_list(doc, [
        "好友申请模块需要处理双向关系去重。用户发送申请前，系统会查询 requester_id 和 receiver_id 两个方向是否已有记录。如果已有 pending 记录，则提示已发送过申请；如果已有 accepted 记录，则提示已经是好友；如果原记录为 rejected，则更新原记录重新发起申请。这样可以避免同一对用户之间产生多条重复好友关系。",
        "搜索模块为了降低 PostgREST 过滤表达式被特殊字符干扰的风险，会对用户输入进行简单清洗，去除逗号、括号、引号、百分号、下划线等特殊字符后再执行 ilike 查询。搜索结果会排除当前用户，避免用户给自己发送好友申请或查看重复资料。",
        "点赞逻辑会根据当前用户 ID 和互动 scope 判断是否已经点赞。如果已经点赞，则本次操作为取消点赞；如果尚未点赞，则追加点赞对象。前端先更新本地 likedBy 数组，然后调用 add_like 或 remove_like。若 RPC 失败，系统恢复 previousLikedBy 并提示用户重试。评论逻辑同样先追加本地评论，再调用 add_comment；删除评论调用 remove_comment 并在失败时回滚。",
        "通知模块采用分页加载，默认每页 20 条。首次加载时，系统同时统计未读通知数量并写入 Zustand Store。用户打开通知中心后，可以标记单条已读或全部已读；当 Realtime 监听到 notifications 表有新记录插入时，系统重新拉取通知。前端还提供通知偏好设置，不同类型通知可以分别控制站内显示、声音和震动。"
    ])
    screenshot_box(doc, "图5-6  动态详情与评论页面真实截图（请替换）")
    screenshot_box(doc, "图5-7  个人中心页面真实截图（请替换）")
    add_title(doc, "5.6  移动端落地实现", 15, level=2)
    add_paragraph(doc, "系统通过 Capacitor 将 Vite 构建后的 dist 目录同步到 Android 和 iOS 原生工程。Android 包名配置为 com.ycy.habit，应用名称为 Habit。移动端可读取原生 App 版本信息，并通过 app_releases 表检查是否存在更高 build_number 的版本。项目中同时保留了 Codemagic 和 Appflow 相关配置，便于后续进行云端构建和发布。")
    add_para_list(doc, [
        "移动端适配中，Vite 的 base 设置为 './'，保证构建产物在 Capacitor 容器中可以通过相对路径正确加载。Capacitor 配置中的 webDir 指向 dist，appId 为 com.ycy.habit，appName 为 Habit。构建流程为先执行 npm run build，再执行 npx cap sync android，同步完成后可通过 Android Studio 或 Gradle 命令生成 APK。",
        "系统还实现了 App 更新检查逻辑。在 Web 预览环境下，当前版本显示为 Web Preview；在原生环境下，系统通过 @capacitor/app 获取当前 version 和 build，再从 Supabase app_releases 表读取最新发布记录。如果最新 build_number 大于当前构建号，则将更新信息写入应用状态，后续可由界面展示更新提示。",
        "本地提醒方面，系统允许用户配置每日提醒开关和提醒时间。提醒配置保存在本地状态中，适合不需要云同步的个性化偏好。通知声音和震动偏好也使用 localStorage 保存，避免频繁写入数据库。这样的设计区分了“需要多端一致的业务数据”和“仅影响本机体验的偏好数据”。"
    ])

    add_title(doc, "第6章  系统测试与部署", 16, level=1, center=True)
    add_title(doc, "6.1  测试环境", 15, level=2)
    add_table(doc, "表6-1  测试环境说明",
              ["项目", "说明"],
              [
                  ["前端运行环境", "Node.js、npm、Vite 开发服务器"],
                  ["开发语言", "TypeScript、TSX、CSS"],
                  ["后端服务", "Supabase Auth、Database、Storage、Realtime"],
                  ["移动端环境", "Android Studio、Capacitor Android 工程"],
                  ["测试方式", "功能测试、界面操作测试、异常流程测试"]
              ], [4.0, 10.0])
    add_para_list(doc, [
        "测试以真实运行的项目为对象，而不是仅对静态页面进行展示。前端在 Vite 开发服务器中进行功能调试，后端连接 Supabase 项目进行真实认证、数据读写和文件上传，移动端通过 Capacitor 同步后在 Android 工程中验证构建流程。测试过程中重点关注数据是否正确写入、界面状态是否及时更新、异常场景是否给出可理解提示。",
        "由于本系统包含多用户协作和社交关系，部分测试需要使用两个或多个账号完成。例如好友申请需要一个账号发送、另一个账号接收；团队挑战需要队长创建邀请码、成员加入并共同打卡；点赞评论需要其他用户对动态进行操作并观察通知是否生成。这类测试能够验证系统是否真正支持社交场景，而不仅是单用户自测。"
    ])
    add_title(doc, "6.2  功能测试", 15, level=2)
    add_table(doc, "表6-2  主要功能测试用例",
              ["编号", "测试项", "操作步骤", "预期结果", "结果"],
              [
                  ["1", "用户登录", "输入已注册邮箱和密码登录", "进入主界面并加载用户资料", "通过"],
                  ["2", "创建习惯", "输入任务名称和目标天数后创建", "任务出现在进行中列表", "通过"],
                  ["3", "每日打卡", "选择任务并发布打卡内容", "生成打卡记录和动态", "通过"],
                  ["4", "重复打卡", "同一天再次打卡同一任务", "提示今日已打卡或更新当日动态", "通过"],
                  ["5", "团队加入", "输入有效邀请码加入团队", "成员列表增加当前用户", "通过"],
                  ["6", "点赞评论", "对动态进行点赞和评论", "动态互动状态更新并生成通知", "通过"],
                  ["7", "头像上传", "选择图片上传头像", "头像更新并同步到用户资料", "通过"],
                  ["8", "移动端构建", "执行构建并同步 Android", "生成可运行原生工程", "通过"]
              ], [1.3, 2.5, 4.5, 4.2, 1.5])
    add_title(doc, "6.2.1  异常流程测试", 14, level=3)
    add_table(doc, "表6-3  异常流程测试用例",
              ["编号", "测试项", "输入或操作", "预期结果", "结果"],
              [
                  ["1", "空任务名", "创建任务时不输入名称", "不创建任务并保持弹窗状态", "通过"],
                  ["2", "无效邀请码", "输入不存在的邀请码", "提示邀请码无效", "通过"],
                  ["3", "重复好友申请", "对同一用户再次发送申请", "提示已发送过申请或已经是好友", "通过"],
                  ["4", "非本人编辑动态", "尝试编辑他人动态", "提示只能修改自己的动态", "通过"],
                  ["5", "非队长删除团队", "成员删除进行中团队任务", "提示只有队长可以删除", "通过"],
                  ["6", "重复打卡", "同一任务同一天再次提交", "不重复增加进度", "通过"]
              ], [1.3, 2.6, 4.0, 4.0, 1.5])
    add_title(doc, "6.2.2  数据一致性测试", 14, level=3)
    add_para_list(doc, [
        "数据一致性测试主要验证前端状态、数据库记录和动态展示是否一致。创建任务后，任务应立即出现在本地列表中，并能在刷新后从 Supabase 重新读取；打卡后，habit_logs 表应新增当天记录，habits 表的进度或成员 lastCheckDate 应更新，activities 表应生成对应动态；删除动态后，动态列表应移除该记录，数据库中对应行也应被删除。",
        "点赞评论测试中，首先观察乐观更新是否立即生效，然后刷新页面确认数据仍存在。若模拟 RPC 失败，前端应回滚到原状态。头像和昵称修改后，个人中心资料应更新，动态列表中自己已发布动态的用户快照也应同步更新。这些测试说明系统不仅能完成单次写入，还能维护多个展示位置之间的数据一致。"
    ])
    add_title(doc, "6.2.3  兼容性与易用性测试", 14, level=3)
    add_para_list(doc, [
        "兼容性测试主要在桌面浏览器预览和 Android 容器中进行。浏览器环境用于快速验证界面和数据逻辑，Android 环境用于验证移动端触摸、页面滚动、图片选择、文件上传、底部导航和原生 App 信息读取。由于系统采用响应式最大宽度布局，页面在手机屏幕中保持类似原生应用的纵向使用体验。",
        "易用性测试关注用户是否能够在不阅读说明的情况下完成核心流程。测试结果表明，用户通常可以通过底部导航找到任务列表和个人中心，通过顶部加号进入创建任务，通过底部打卡入口发布记录。对于团队、设置和通知等低频功能，系统使用标签、图标和弹窗降低操作复杂度。"
    ])
    add_title(doc, "6.3  移动端构建与部署", 15, level=2)
    add_paragraph(doc, "系统构建过程包括 npm install 安装依赖、npm run build 生成 Web 构建产物、npx cap sync android 同步到 Android 工程以及使用 Gradle 生成 APK。由于项目已经落地，后续提交论文时可在本章补充真实 APK 安装截图、手机运行截图或云构建记录截图。")
    add_para_list(doc, [
        "实际部署时，首先需要确认 .env 中已经配置 VITE_SUPABASE_URL 和 VITE_SUPABASE_ANON_KEY，否则前端虽然可以启动，但无法完成真实认证和数据库访问。其次需要确认 Supabase 数据库表、RLS 策略、Storage bucket 和 RPC 函数已经创建，否则部分功能会出现创建失败、点赞失败或图片上传失败。",
        "Web 构建完成后，dist 目录即为静态资源产物。执行 npx cap sync android 后，Capacitor 会将 dist 同步到 android/app/src/main/assets/public 等相关目录，并更新原生插件配置。随后可以在 Android Studio 中运行，也可以进入 android 目录执行 .\\gradlew.bat assembleDebug 生成 Debug APK。生成的 APK 可安装到手机上进行真实测试。",
        "云构建方面，项目保留了 codemagic.yaml 和 CODEMAGIC_GUIDE.md，用于后续配置 Codemagic 构建。云构建的价值在于可以在相对稳定的环境中完成依赖安装、Web 构建、Capacitor 同步和 Android 打包，并保存构建产物。若后续进入正式发布阶段，还需要进一步配置签名证书、版本号、隐私政策和应用商店素材。"
    ])
    add_title(doc, "6.4  测试结果分析", 15, level=2)
    add_paragraph(doc, "测试结果表明，系统能够完成毕业设计要求中的主要业务功能，界面流程较为完整，用户可以完成从注册登录、创建任务、每日打卡到社交互动和团队挑战的闭环操作。测试过程中也发现，部分复杂业务如断签检查、团队投票超时等目前主要依赖客户端触发，后续可通过服务端定时任务进一步提高可靠性。")
    add_para_list(doc, [
        "从功能覆盖角度看，系统已经覆盖习惯养成应用的核心路径：用户能够注册账号、维护资料、创建任务、连续打卡、查看进度、发布动态、获得勋章，并通过好友、关注、评论和通知获得反馈。团队挑战功能进一步扩展了单人习惯场景，使作品具有一定差异化。",
        "从工程质量角度看，项目已经具备较清晰的前端模块划分和云端数据结构，关键操作具有失败提示和部分回滚能力。但由于开发周期有限，系统仍存在一些可优化点：第一，部分 RLS 策略需要进一步细化，特别是 activities 表的互动更新权限；第二，断签检查和投票超时更适合在服务端定时执行；第三，自动化测试覆盖不足，目前主要依赖人工测试；第四，团队成员 JSONB 方案后续在复杂统计和权限控制上会受到限制。",
        "总体而言，测试结果证明该系统可以作为毕业设计作品完整展示，也具备进一步产品化的基础。后续若继续开发，应优先补齐数据库 migration、服务端定时任务、自动化测试和正式发布流程。"
    ])

    add_title(doc, "结论", 18, center=True, before=12)
    for t in [
        "本文围绕 Habit 习惯追踪社交 App 的设计与实现展开，完成了需求分析、总体架构设计、数据库设计、功能模块实现、测试与部署说明等工作。系统以 React、TypeScript、Vite 和 Tailwind CSS 为前端基础，以 Supabase 为云后端，以 Capacitor 为移动端封装方案，实现了一个可运行、可扩展、具有实际使用价值的毕业设计作品。",
        "从功能实现结果看，系统已经支持用户认证、个人资料、习惯创建、每日打卡、图片上传、动态广场、好友关注、点赞评论、通知中心、团队挑战和勋章奖励等功能，能够形成较完整的用户使用闭环。从工程实践看，项目将前端组件化、状态管理、后端数据建模、实时订阅、移动端打包和云服务集成结合起来，提高了本人对现代移动应用开发流程的理解。",
        "当然，系统仍存在进一步优化空间。例如，团队成员数据可从 JSONB 逐步迁移为关系表，断签检查和投票超时处理可迁移到服务端定时任务，RLS 策略和 RPC 权限可继续细化，测试体系也可以补充自动化测试。后续如果继续迭代，可围绕数据统计、习惯推荐、更多提醒策略和正式应用商店发布进行完善。"
    ]:
        add_paragraph(doc, t)

    add_title(doc, "参考文献", 18, center=True, before=12)
    refs = [
        "[1] Facebook Open Source. React Documentation[EB/OL]. https://react.dev/.",
        "[2] Microsoft. TypeScript Handbook[EB/OL]. https://www.typescriptlang.org/docs/.",
        "[3] Vite Team. Vite Documentation[EB/OL]. https://vite.dev/.",
        "[4] Supabase. Supabase Documentation[EB/OL]. https://supabase.com/docs/.",
        "[5] Ionic Team. Capacitor Documentation[EB/OL]. https://capacitorjs.com/docs/.",
        "[6] PostgreSQL Global Development Group. PostgreSQL Documentation[EB/OL]. https://www.postgresql.org/docs/.",
        "[7] Tailwind Labs. Tailwind CSS Documentation[EB/OL]. https://tailwindcss.com/docs/.",
        "[8] Gamma E,Helm R,Johnson R,Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Boston:Addison-Wesley,1994.",
        "[9] Pressman R S. Software Engineering: A Practitioner's Approach[M]. New York:McGraw-Hill,2014.",
        "[10] Sommerville I. Software Engineering[M]. Boston:Pearson,2016."
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_indent=False)

    add_title(doc, "附录", 18, center=True, before=12)
    add_title(doc, "附录A  主要运行命令", 15, level=2)
    for cmd in [
        "npm install：安装项目依赖。",
        "npm run dev：启动本地开发服务器，默认端口为 3000。",
        "npm run build：构建生产环境 Web 资源。",
        "npx cap sync android：将 Web 构建产物同步到 Android 工程。",
        ".\\gradlew.bat assembleDebug：在 Windows 环境中构建 Android Debug APK。"
    ]:
        add_paragraph(doc, cmd)
    add_title(doc, "附录B  截图替换清单", 15, level=2)
    for item in [
        "图5-1：首页/广场页面截图。",
        "图5-2：任务列表页面截图。",
        "图5-3：登录或注册页面截图。",
        "图5-4：打卡发布页面截图。",
        "图5-5：团队挑战页面截图。",
        "图5-6：动态详情与评论页面截图。",
        "图5-7：个人中心页面截图。"
    ]:
        add_paragraph(doc, item)

    add_title(doc, "致谢", 18, center=True, before=12)
    for t in [
        "本毕业设计从选题、需求分析、系统开发到论文撰写，得到了老师和同学的帮助。在此，首先感谢指导教师在毕业设计过程中给予的指导和建议，使我能够不断明确系统目标、完善功能设计并规范论文结构。",
        "同时感谢同学和朋友在应用测试过程中提供的反馈，帮助我发现界面交互和业务流程中的不足。通过本次毕业设计，我对前端工程化、云后端服务、移动端打包和完整项目落地有了更加系统的认识，也提升了解决实际问题的能力。"
    ]:
        add_paragraph(doc, t)


def finalize_styles(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("第") and "章" in text:
            for r in p.runs:
                set_run_font(r, "黑体", 16, True)
        elif text and len(text) > 2 and text[0].isdigit() and "." in text[:5]:
            parts = text.split()[0]
            dots = parts.count(".")
            for r in p.runs:
                set_run_font(r, "黑体", 15 if dots == 1 else 14 if dots == 2 else 12, True)


def main():
    doc = Document()
    set_doc_defaults(doc)
    set_section_page(doc.sections[0])
    set_headers_footers(doc.sections[0], False)

    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)
    add_body(doc)
    finalize_styles(doc)
    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        doc.save(TMP_OUT)
        print(TMP_OUT)


if __name__ == "__main__":
    main()
